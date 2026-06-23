import docx
import os

# Use glob to find the file regardless of encoding
docs_dir = r'C:\Users\SuperUsuario\Documents\Proyectos\Servi_Elec_Manager\docs'
files = os.listdir(docs_dir)
print("Files found:")
for f in files:
    print(f"  {repr(f)}")

base_file = [f for f in files if '2.1.1' in f][0]
base_path = os.path.join(docs_dir, base_file)
print(f"\nOpening: {base_path}")

doc = docx.Document(base_path)
print(f'Total parrafos: {len(doc.paragraphs)}')
print(f'Total tablas: {len(doc.tables)}')

# Count images
from docx.oxml.ns import qn
body = doc.element.body
imgs = body.findall('.//' + qn('a:blip'))
print(f'Referencias imagen (blip): {len(imgs)}')

print('\n--- TODOS LOS PARRAFOS CON ESTILO (primeros 300) ---')
for i, p in enumerate(doc.paragraphs[:300]):
    txt = p.text.strip()
    style = p.style.name
    if txt:
        print(f'[{style}] {txt[:150]}')
