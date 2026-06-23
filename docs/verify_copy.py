import zipfile
from docx import Document

src = "2.1.1 Documento de Informe Implementación del Proyecto (1).docx"
dst = "3.1.1 Documento de Informe Implementación del Proyecto.docx"

def inspect(path):
    doc = Document(path)
    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)
    with zipfile.ZipFile(path) as z:
        images = [f for f in z.namelist() if f.startswith("word/media/")]
    return paragraphs, tables, len(images)

p2, t2, i2 = inspect(src)
p3, t3, i3 = inspect(dst)

print(f"{'':30} {'DOC 2 (BASE)':>15} {'DOC 3 (DEST)':>15} {'IGUAL?':>8}")
print("-" * 72)
print(f"{'Parrafos':30} {p2:>15} {p3:>15} {'OK' if p2==p3 else 'DIFF':>8}")
print(f"{'Tablas':30} {t2:>15} {t3:>15} {'OK' if t2==t3 else 'DIFF':>8}")
print(f"{'Imagenes (word/media/)':30} {i2:>15} {i3:>15} {'OK' if i2==i3 else 'DIFF':>8}")
print()
if p2==p3 and t2==t3 and i2==i3:
    print("VERIFICACION COMPLETA: doc 3 es identico al doc 2.")
else:
    print("DIFERENCIAS DETECTADAS.")
