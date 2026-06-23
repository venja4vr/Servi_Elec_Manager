from docx import Document

doc = Document("3.1.1 Documento de Informe Implementación del Proyecto.docx")

# Buscar secciones Glosario y Bibliografía
print("=== BUSCANDO GLOSARIO Y BIBLIOGRAFÍA ===\n")

glosario_idx = None
biblio_idx = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    text_upper = text.upper()
    if "GLOSARIO" in text_upper:
        print(f"[{i:4d}] GLOSARIO encontrado | estilo='{p.style.name}' | texto='{text[:80]}'")
        if glosario_idx is None:
            glosario_idx = i
    if "BIBLIOGRAF" in text_upper:
        print(f"[{i:4d}] BIBLIOGRAFÍA encontrada | estilo='{p.style.name}' | texto='{text[:80]}'")
        if biblio_idx is None:
            biblio_idx = i

print()

# Mostrar el bloque de Glosario (hasta 60 párrafos)
if glosario_idx is not None:
    print(f"=== CONTENIDO GLOSARIO (desde idx {glosario_idx}) ===\n")
    end = biblio_idx if biblio_idx else glosario_idx + 60
    for i in range(glosario_idx, min(end, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if text:
            print(f"[{i:4d}] estilo='{p.style.name}' | '{text[:100]}'")
            # Mostrar runs para ver formato
            for j, run in enumerate(p.runs):
                if run.text.strip():
                    print(f"       run[{j}] bold={run.bold} | '{run.text[:80]}'")

print()

# Mostrar el bloque de Bibliografía (hasta 30 párrafos)
if biblio_idx is not None:
    print(f"=== CONTENIDO BIBLIOGRAFÍA (desde idx {biblio_idx}) ===\n")
    for i in range(biblio_idx, min(biblio_idx + 30, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if text:
            print(f"[{i:4d}] estilo='{p.style.name}' | '{text[:100]}'")

print(f"\nTotal párrafos doc: {len(doc.paragraphs)}")
