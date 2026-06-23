import docx
import os

docs_dir = r'C:\Users\SuperUsuario\Documents\Proyectos\Servi_Elec_Manager\docs'
files = os.listdir(docs_dir)
base_file = [f for f in files if '2.1.1' in f][0]
base_path = os.path.join(docs_dir, base_file)

doc = docx.Document(base_path)

# Print ALL paragraphs with styles (full document)
print(f'=== DOCUMENT STATS ===')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')
from docx.oxml.ns import qn
imgs = doc.element.body.findall('.//' + qn('a:blip'))
print(f'Images (blip refs): {len(imgs)}')

print('\n=== ALL HEADINGS & STRUCTURE ===')
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    style = p.style.name
    if not txt:
        continue
    if any(x in style for x in ['Heading', 'tulo', 'Title', 'TOC']):
        level = '  ' * (int(style[-1]) - 1) if style[-1].isdigit() else ''
        print(f'{level}[{style}] {txt[:200]}')
    elif style in ['table of figures']:
        print(f'  [TOF] {txt[:150]}')

print('\n=== ALL PARAGRAPHS (para detectar secciones futuras) ===')
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    style = p.style.name
    if txt and len(txt) > 20:
        print(f'[{i}][{style}] {txt[:200]}')
