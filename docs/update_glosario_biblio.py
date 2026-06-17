# encoding: utf-8
"""
Chunk 2: Inserta 8 términos en Glosario y 4 fuentes en Bibliografía del doc 3.
Usa referencias de elementos XML (no índices) para no invalidar posiciones al insertar.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy

DOC_PATH = "3.1.1 Documento de Informe Implementación del Proyecto.docx"

doc = Document(DOC_PATH)
paras = doc.paragraphs
n_before = len(paras)
print(f"Párrafos antes: {n_before}")

# ── Verificar que encontramos las secciones correctas ──────────────────────────
glos_idx = next(i for i, p in enumerate(paras) if "GLOSARIO" in p.text.upper() and p.style.name.startswith("Heading"))
bib_idx  = next(i for i, p in enumerate(paras) if "BIBLIOGRAF" in p.text.upper() and p.style.name.startswith("Heading"))
print(f"GLOSARIO en idx {glos_idx}, BIBLIOGRAFÍA en idx {bib_idx}")
print(f"Estilo del primer término del glosario: '{paras[glos_idx+1].style.name}'")
print(f"Estilo de entrada bibliografía ejemplo: '{paras[bib_idx+2].style.name}'")

# ── Recolectar referencias XML ANTES de cualquier modificación ─────────────────
# Glosario: insertar ANTES de los siguientes párrafos
ref_base_datos  = paras[151]._element   # Base de Datos → Badge dinámico va antes
ref_historia    = paras[156]._element   # Historia de Usuario → Groq va antes
ref_mer         = paras[158]._element   # MER → LLM va antes
ref_postgresql  = paras[160]._element   # PostgreSQL → Notificación y Outlier van antes
ref_vite        = paras[168]._element   # Vite → Throttling va antes
ref_wireframe   = paras[169]._element   # Wireframe → VTEX va antes (después VTEX)
                                        # Zonificación va DESPUÉS de Wireframe

# Bibliografía: insertar DESPUÉS del último entry
ref_biblio_last = paras[1043]._element  # ISO/IEC/IEEE — última entrada

# Plantillas para clonar estilos
tpl_list_para   = paras[149]._element   # Lista de glosario (List Paragraph)
tpl_biblio_cat  = paras[bib_idx + 1]._element  # Encabezado de categoría biblio (Normal)
tpl_biblio_entry = paras[bib_idx + 2]._element  # Entrada biblio (List Paragraph)

print(f"Plantillas OK")


# ── Helpers ────────────────────────────────────────────────────────────────────

def _strip_runs(elem):
    """Elimina todos los hijos de w:p excepto w:pPr."""
    to_remove = [ch for ch in elem if etree.QName(ch).localname != 'pPr']
    for ch in to_remove:
        elem.remove(ch)


def _add_run(p_elem, text, bold=False):
    r = OxmlElement('w:r')
    if bold:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p_elem.append(r)


def make_glosario_entry(term_bold, definition_normal):
    """Clona plantilla List Paragraph y pone term en bold + definición normal."""
    new_p = copy.deepcopy(tpl_list_para)
    _strip_runs(new_p)
    _add_run(new_p, term_bold, bold=True)
    _add_run(new_p, definition_normal, bold=False)
    return new_p


def make_biblio_category(text):
    """Clona encabezado de categoría (Normal) con texto plano."""
    new_p = copy.deepcopy(tpl_biblio_cat)
    _strip_runs(new_p)
    _add_run(new_p, text)
    return new_p


def make_biblio_entry(text):
    """Clona entrada bibliográfica (List Paragraph) con texto plano."""
    new_p = copy.deepcopy(tpl_biblio_entry)
    _strip_runs(new_p)
    _add_run(new_p, text)
    return new_p


# ── PASO 1: Insertar términos en Glosario ─────────────────────────────────────
print("\n--- Insertando en Glosario ---")

# 1. Badge dinámico  (B-ad → después de Backend B-ac, antes de Base de Datos B-as)
ref_base_datos.addprevious(make_glosario_entry(
    "Badge dinámico:",
    " Indicador visual numérico en la interfaz que muestra la cantidad de notificaciones "
    "no leídas, cambiando de color (verde a naranja) cuando se detectan eventos nuevos "
    "desde la última visita."
))
print("  ✓ Badge dinámico")

# 2. Groq  (G → después de Frontend F, antes de Historia H)
ref_historia.addprevious(make_glosario_entry(
    "Groq:",
    " Servicio en la nube que ofrece inferencia rápida de modelos de lenguaje grandes "
    "(LLMs) open source como Llama 3.3 mediante una API compatible con OpenAI. "
    "Permite ejecutar IA sin costos en el tier gratuito."
))
print("  ✓ Groq")

# 3. LLM  (L → después de JWT J, antes de MER M)
ref_mer.addprevious(make_glosario_entry(
    "LLM (Large Language Model):",
    " Modelo de lenguaje grande entrenado con millones de textos, capaz de comprender y "
    "generar lenguaje natural. En este proyecto se utiliza Llama 3.3 70B vía Groq."
))
print("  ✓ LLM")

# 4. Notificación in-app  (N → después de Microservicios M, antes de Outlier O, antes de PostgreSQL P)
ref_postgresql.addprevious(make_glosario_entry(
    "Notificación in-app:",
    " Aviso o alerta que se muestra dentro de la aplicación (no por correo o WhatsApp), "
    "accesible mediante una sección dedicada y un badge en el menú lateral."
))
print("  ✓ Notificación in-app")

# 5. Outlier de precio  (O → antes de PostgreSQL P, después de Notificación N ya insertada)
ref_postgresql.addprevious(make_glosario_entry(
    "Outlier de precio:",
    " Valor anómalo en un histórico de precios que se aparta significativamente de la "
    "mediana de los registros válidos. El sistema los detecta automáticamente para no "
    "contaminar el promedio mostrado."
))
print("  ✓ Outlier de precio")

# 6. Throttling  (T → después de Swagger UI Sw, antes de Vite V)
ref_vite.addprevious(make_glosario_entry(
    "Throttling:",
    " Técnica que limita la frecuencia con que se generan eventos (en este proyecto, "
    "notificaciones) para evitar acumulación excesiva. "
    "Se usa una ventana de 5 segundos por tipo de evento."
))
print("  ✓ Throttling")

# 7. VTEX  (VT → después de Vite Vi, antes de Wireframe W)
ref_wireframe.addprevious(make_glosario_entry(
    "VTEX:",
    " Plataforma de comercio electrónico utilizada por Easy y Sodimac. Su API pública "
    "permite obtener precios y disponibilidad en tiempo real sin scraping web tradicional."
))
print("  ✓ VTEX")

# 8. Zonificación  (Z → DESPUÉS de Wireframe, último del glosario)
ref_wireframe.addnext(make_glosario_entry(
    "Zonificación:",
    " Agrupación de comunas por proximidad geográfica para facilitar el cálculo de costos "
    "de bencina y traslado. En el proyecto se definen 8 zonas que cubren 32 comunas de "
    "la Quinta Región."
))
print("  ✓ Zonificación")

# ── PASO 2: Insertar fuentes en Bibliografía ──────────────────────────────────
print("\n--- Insertando en Bibliografía ---")

# Agregar categoría + 4 entradas después del último entry de bibliografía
p_cat = make_biblio_category("AI, Machine Learning & E-commerce APIs")
ref_biblio_last.addnext(p_cat)

last = p_cat
for text in [
    "Groq, Inc. (2024). Groq API Documentation. https://console.groq.com/docs",
    "OpenAI. (2024). Python SDK for OpenAI API. https://github.com/openai/openai-python",
    "Meta AI. (2024). Llama 3.3: 70B Instruction-Tuned Model. https://ai.meta.com/llama/",
    "VTEX. (2024). VTEX IO API Reference. https://developers.vtex.com/",
]:
    p = make_biblio_entry(text)
    last.addnext(p)
    last = p

print("  ✓ 4 fuentes + categoría AI/ML/Ecommerce")

# ── PASO 3: Guardar ───────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print(f"\n✓ Documento guardado: {DOC_PATH}")

# ── PASO 4: Verificar ─────────────────────────────────────────────────────────
doc2 = Document(DOC_PATH)
n_after = len(doc2.paragraphs)

print(f"\n{'─'*60}")
print(f"Párrafos antes:  {n_before}")
print(f"Párrafos después:{n_after}")
print(f"Diferencia:      +{n_after - n_before}  (esperado +13: 8 glosario + 5 biblio)")
print(f"{'─'*60}")

print("\n=== VERIFICACIÓN: términos nuevos en Glosario ===")
GLOS_TERMS = [
    "Badge dinámico", "Groq:", "LLM (Large Language",
    "Notificación in-app", "Outlier de precio",
    "Throttling:", "VTEX:", "Zonificación:"
]
found_glos = 0
for i, p in enumerate(doc2.paragraphs):
    if any(t in p.text for t in GLOS_TERMS):
        print(f"  [{i:4d}] {p.text[:90]}")
        found_glos += 1
print(f"  → {found_glos}/8 términos encontrados")

print("\n=== VERIFICACIÓN: fuentes nuevas en Bibliografía ===")
BIB_TERMS = ["Groq, Inc.", "OpenAI. (2024)", "Meta AI.", "VTEX. (2024)", "AI, Machine"]
found_bib = 0
for i, p in enumerate(doc2.paragraphs):
    if any(t in p.text for t in BIB_TERMS):
        print(f"  [{i:4d}] {p.text[:90]}")
        found_bib += 1
print(f"  → {found_bib}/5 entradas encontradas (4 fuentes + 1 categoría)")
